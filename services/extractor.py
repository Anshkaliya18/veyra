# services/extractor.py
from __future__ import annotations

import csv
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Optional

import pandas as pd

logger = logging.getLogger(__name__)

# Optional imports — these libraries are only needed for the file types they support.
try:
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover
    fitz = None

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None

try:
    from PIL import Image
except Exception:  # pragma: no cover
    Image = None

try:
    import pytesseract
except Exception:  # pragma: no cover
    pytesseract = None


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".csv",
    ".xlsx",
    ".xls",
    ".png",
    ".jpg",
    ".jpeg",
}

# Note:
# - ".doc" is not supported here because it needs special conversion tooling.
# - If you want ".doc" support later, convert to .docx before extraction or
#   add a LibreOffice/textract-based fallback.


@dataclass
class ExtractionResult:
    success: bool
    text: str
    file_type: str
    pages: int = 0
    error: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        return {
            "success": self.success,
            "text": self.text,
            "file_type": self.file_type,
            "pages": self.pages,
            "error": self.error,
        }


class ExtractionError(Exception):
    pass


def extract_text(file_path: str) -> Dict[str, Any]:
    """
    Main entry point.

    Returns:
        {
            "success": bool,
            "text": str,
            "file_type": str,
            "pages": int,
            "error": str | None
        }
    """
    path = Path(file_path)

    if not path.exists():
        return ExtractionResult(
            success=False,
            text="",
            file_type="unknown",
            pages=0,
            error=f"File not found: {file_path}",
        ).to_dict()

    ext = path.suffix.lower().strip()

    if ext not in SUPPORTED_EXTENSIONS:
        return ExtractionResult(
            success=False,
            text="",
            file_type=ext.lstrip(".") or "unknown",
            pages=0,
            error=f"Unsupported file type: {ext}",
        ).to_dict()

    try:
        if ext == ".pdf":
            text, pages = _extract_pdf(path)
            return ExtractionResult(True, text, "pdf", pages).to_dict()

        if ext == ".docx":
            text = _extract_docx(path)
            return ExtractionResult(True, text, "docx", 0).to_dict()

        if ext == ".txt":
            text = _extract_txt(path)
            return ExtractionResult(True, text, "txt", 0).to_dict()

        if ext == ".csv":
            text = _extract_csv(path)
            return ExtractionResult(True, text, "csv", 0).to_dict()

        if ext in {".xlsx", ".xls"}:
            text = _extract_excel(path)
            return ExtractionResult(True, text, "excel", 0).to_dict()

        if ext in {".png", ".jpg", ".jpeg"}:
            text = _extract_image(path)
            return ExtractionResult(True, text, "image", 0).to_dict()

        return ExtractionResult(
            success=False,
            text="",
            file_type=ext.lstrip("."),
            pages=0,
            error=f"No extractor implemented for: {ext}",
        ).to_dict()

    except Exception as exc:
        logger.exception("Text extraction failed for %s", file_path)
        return ExtractionResult(
            success=False,
            text="",
            file_type=ext.lstrip(".") or "unknown",
            pages=0,
            error=str(exc),
        ).to_dict()


def _clean_text(text: str) -> str:
    if not text:
        return ""
    # Keep paragraphs readable, remove excessive blank lines/spaces.
    text = text.replace("\x00", "")
    lines = [line.rstrip() for line in text.splitlines()]
    cleaned = "\n".join(lines)
    while "\n\n\n" in cleaned:
        cleaned = cleaned.replace("\n\n\n", "\n\n")
    return cleaned.strip()

def _extract_pdf(path: Path) -> tuple[str, int]:
    if fitz is None:
        raise ExtractionError(
            "PyMuPDF is not installed. Run: pip install pymupdf"
        )

    if Image is None:
        raise ExtractionError(
            "Pillow is not installed. Run: pip install pillow"
        )

    if pytesseract is None:
        raise ExtractionError(
            "pytesseract is not installed. Run: pip install pytesseract"
        )

    pages_text = []
    pages_count = 0

    with fitz.open(str(path)) as pdf:
        pages_count = len(pdf)

        for page in pdf:

            # Try extracting embedded PDF text first
            page_text = page.get_text("text").strip()

            # If no text exists, use OCR
            if not page_text:

                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))

                img = Image.frombytes(
                    "RGB",
                    [pix.width, pix.height],
                    pix.samples
                )

                page_text = pytesseract.image_to_string(img)

            page_text = _clean_text(page_text)

            if page_text:
                pages_text.append(page_text)

    text = "\n\n".join(pages_text)

    return text, pages_count

def _extract_docx(path: Path) -> str:
    if Document is None:
        raise ExtractionError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    doc = Document(str(path))
    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = (cell.text or "").strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                parts.append(" | ".join(row_text))

    return _clean_text("\n".join(parts))


def _extract_txt(path: Path) -> str:
    # Try utf-8 first, then fall back to a forgiving read.
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="utf-8", errors="ignore")
    return _clean_text(text)


def _extract_csv(path: Path) -> str:
    # pandas is the easiest reliable choice for CSV to readable text.
    df = pd.read_csv(path)
    # Keep it compact but readable.
    return _clean_text(df.to_string(index=False))


def _extract_excel(path: Path) -> str:
    # Read all sheets if possible.
    sheets = pd.read_excel(path, sheet_name=None)
    parts: list[str] = []

    for sheet_name, df in sheets.items():
        parts.append(f"Sheet: {sheet_name}")
        parts.append(df.to_string(index=False))
        parts.append("")

    return _clean_text("\n".join(parts))


def _extract_image(path: Path) -> str:
    if Image is None:
        raise ExtractionError(
            "Pillow is not installed. Run: pip install pillow"
        )
    if pytesseract is None:
        raise ExtractionError(
            "pytesseract is not installed. Run: pip install pytesseract"
        )

    try:
        img = Image.open(path)
        text = pytesseract.image_to_string(img)
        return _clean_text(text)
    except Exception as exc:
        raise ExtractionError(f"OCR failed: {exc}") from exc
    